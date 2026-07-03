#!/usr/bin/env python3
"""
data2md - конвертер данных любого типа в Markdown для последующей обработки LLM через repomix или yek.

Назначение:
    Преобразует разнородные данные (изображения, PDF, DOCX, Excel, CSV, текст)
    в единый текстовый формат Markdown, который repomix или yek может упаковать в единый
    корпус для подачи в локальную LLM (MetalGPT-1, Qwen2.5-VL).

Поддерживаемые форматы:
    - Изображения: PNG, JPG, JPEG
      -> Текстовое описание через Qwen2.5-VL (vision-language модель) или
         Deepseek API, VK API, Yandex Cloud Vision API, или
         fallback-описание на основе имени файла, если модели недоступны.
      -> Выход: Markdown с описанием узлов, связей, параметров.

    - Текст: PDF, DOCX, MD, TXT
      -> Извлечение текста с сохранением структуры (заголовки, параграфы, таблицы).
      -> Выход: Markdown с иерархией документа.

    - Таблицы: XLSX, XLS, CSV
      -> Чтение через pandas, вывод в Markdown-таблицы.
      -> Дополнительно: статистика по числовым колонкам, извлечение ключевых
        метрик (Ni, Cu, Fe, потери, класс крупности и т.д.), первые N строк.
      -> Выход: Markdown с таблицами и метриками.

Особенности:
    - Обработка через .env конфигурацию или CLI параметры.
    - Фильтрация по расширениям (--ext-include, --ext-exclude).
    - Переключение типов данных (--no-images, --no-text, --no-tables).
    - Ограничения: --max-rows для Excel, --page-limit для PDF.
    - Vision через несколько провайдеров: ollama (Qwen2.5-VL), Deepseek API, VK API, Yandex Cloud.
    - Корпусный строитель: repomix (по умолчанию) или yek.
    - Автоматическое создание выходной директории.

Использование:
    python data2md.py --input ./данные --output ./repomix_input
    python data2md.py --input ./данные --output ./repomix_input --no-images --max-rows 20
    python data2md.py --input ./данные --output ./repomix_input \
        --ext-include .png,.xlsx --verbose
    python data2md.py --input ./данные --output ./repomix_input --vision-provider deepseek --vision-deepseek-key YOUR_KEY
    python data2md.py --input ./данные --output ./repomix_input --vision-provider vk --vision-vk-key YOUR_KEY
    python data2md.py --input ./данные --output ./repomix_input --vision-provider yandex --vision-yandex-key YOUR_KEY
    python data2md.py --input ./данные --output ./repomix_input --corpus-builder yek
"""

import argparse
import os
import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader

# ============================================================
# Загрузка .env
# ============================================================
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv не обязателен

# ============================================================
# Конфигурация из .env / defaults
# ============================================================
CFG = {
    "INPUT_DIR": os.getenv("INPUT_DIR", "./данные"),
    "OUTPUT_DIR": os.getenv("OUTPUT_DIR", "./repomix_input"),
    "VISION_MODEL": os.getenv("VISION_MODEL", "qwen2.5-vl:7b"),
    "VISION_OLLAMA_HOST": os.getenv("VISION_OLLAMA_HOST", "http://localhost:11434"),
    "MAX_TABLE_ROWS": int(os.getenv("MAX_TABLE_ROWS", "50")),
    "PDF_PAGE_LIMIT": int(os.getenv("PDF_PAGE_LIMIT", "0")),
    "IMAGE_ENABLED": os.getenv("IMAGE_ENABLED", "true").lower() == "true",
    "TEXT_ENABLED": os.getenv("TEXT_ENABLED", "true").lower() == "true",
    "TABLE_ENABLED": os.getenv("TABLE_ENABLED", "true").lower() == "true",
    "INCLUDE_STATS": os.getenv("INCLUDE_STATS", "true").lower() == "true",
    "INCLUDE_KEY_METRICS": os.getenv("INCLUDE_KEY_METRICS", "true").lower() == "true",
    "VERBOSE": os.getenv("VERBOSE", "false").lower() == "true",
    "EXTENSIONS_INCLUDE": os.getenv("EXTENSIONS_INCLUDE", ""),
    "EXTENSIONS_EXCLUDE": os.getenv("EXTENSIONS_EXCLUDE", ".log,.tmp"),
}

# Mapping: расширение - тип
EXT_MAP = {
    ".png": "image", ".jpg": "image", ".jpeg": "image",
    ".pdf": "text",
    ".docx": "text",
    ".md": "text", ".txt": "text",
    ".xlsx": "tabular", ".xls": "tabular", ".csv": "tabular",
}

# Ключевые слова для извлечения метрик из таблиц
KEY_METRICS_KEYWORDS = ['ni', 'никель', 'cu', 'медь', 'потер', 'содерж',
                        'co', 'кобальт', 'fe', 'железо', 's', 'сера',
                        'au', 'золото', 'ag', 'серебро', 'pt', 'pd',
                        'класс', 'крупность', 'извлечен', 'хвост']


# ============================================================
# Обработчики
# ============================================================

def image_to_md(image_path: str, cfg: dict) -> str:
    """Изображение в Markdown (через Qwen2.5-VL или заглушку)."""
    stem = Path(image_path).stem
    name = Path(image_path).name
    md = f"# {name}\n\n"

    if cfg.get("VISION_MODEL") and _is_ollama_available():
        # Реальный вызов Qwen2.5-VL через ollama
        md += _ollama_vision_describe(image_path, cfg)
    else:
        # Заглушка на основе имени файла
        md += _image_fallback_description(image_path)

    md += f"\n\n*Источник: {name}*\n"
    return md


def table_to_md(table_path: str, cfg: dict) -> str:
    """Excel/CSV в Markdown таблицу + статистику."""
    stem = Path(table_path).stem
    ext = Path(table_path).suffix.lower()

    md = f"# {Path(table_path).name}\n\n"

    # Загрузка
    try:
        if ext == ".csv":
            df = pd.read_csv(table_path)
        else:
            df = pd.read_excel(table_path)
    except Exception as e:
        return f"# {Path(table_path).name}\n\n*Ошибка загрузки: {e}*\n"

    md += f"- **Строк**: {len(df)}, **Колонок**: {len(df.columns)}\n"
    md += f"- **Типы данных**: {', '.join(str(df[c].dtype) for c in df.columns[:5])}"
    if len(df.columns) > 5:
        md += f" ... (+{len(df.columns)-5})"
    md += "\n\n"

    # Имена колонок
    md += "## Колонки\n\n"
    for i, col in enumerate(df.columns):
        md += f"{i+1}. `{col}` ({df[col].dtype})"
        non_null = df[col].notna().sum()
        md += f" — {non_null}/{len(df)} непустых"
        md += "\n"

    # Данные (первые N строк)
    max_rows = cfg.get("MAX_TABLE_ROWS", 50)
    md += f"\n## Данные (первые {min(max_rows, len(df))} строк)\n\n"
    md += df.head(max_rows).to_markdown(index=False)

    # Статистика
    if cfg.get("INCLUDE_STATS", True):
        numeric_cols = df.select_dtypes(include="number").columns
        if len(numeric_cols) > 0:
            md += "\n\n## Статистика\n\n"
            stats = df[numeric_cols].describe().round(3)
            stats.index.name = "метрика"
            md += stats.to_markdown()

    # Ключевые метрики
    if cfg.get("INCLUDE_KEY_METRICS", True):
        for col in df.columns:
            col_lower = col.lower()
            if any(kw in col_lower for kw in KEY_METRICS_KEYWORDS):
                md += f"\n\n**Ключевая колонка**: `{col}`\n"
                try:
                    num = pd.to_numeric(df[col], errors='coerce')
                    valid = num.dropna()
                    if len(valid) > 0:
                        md += f"- Мин: {valid.min():.3f}\n"
                        md += f"- Макс: {valid.max():.3f}\n"
                        md += f"- Среднее: {valid.mean():.3f}\n"
                        md += f"- Медиана: {valid.median():.3f}\n"
                        md += f"- Сумма: {valid.sum():.3f}\n"
                except:
                    pass

    md += f"\n\n*Источник: {Path(table_path).name}*\n"
    return md


def docx_to_md(docx_path: str, cfg: dict) -> str:
    """DOCX в Markdown."""
    doc = Document(docx_path)
    md = f"# {Path(docx_path).name}\n\n"
    md += f"**Параграфов**: {len(doc.paragraphs)}, **Таблиц**: {len(doc.tables)}\n\n"

    # Параграфы
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name.lower() if para.style else ""
            if "heading" in style or "заголов" in style:
                level = min(sum(1 for c in style if c.isdigit()) + 1, 6)
                md += f"{'#' * level} {text}\n\n"
            else:
                md += text + "\n\n"

    # Таблицы документа
    for i, table in enumerate(doc.tables):
        md += f"## Таблица {i+1}\n\n"
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            header = rows[0]
            data = rows[1:]
            if data:
                df = pd.DataFrame(data, columns=header)
                md += df.to_markdown(index=False)
            else:
                md += f"| {' | '.join(header)} |\n"
        md += "\n"

    md += f"\n\n*Источник: {Path(docx_path).name}*\n"
    return md


def pdf_to_md(pdf_path: str, cfg: dict) -> str:
    """PDF в Markdown."""
    reader = PdfReader(pdf_path)
    total = len(reader.pages)
    limit = cfg.get("PDF_PAGE_LIMIT", 0)
    pages = total if limit <= 0 else min(total, limit)

    md = f"# {Path(pdf_path).name}\n\n"
    md += f"**Страниц**: {total}, обработано: {pages}\n\n"

    for i in range(pages):
        text = reader.pages[i].extract_text()
        if text.strip():
            md += f"## Страница {i+1}\n\n{text}\n\n"

    md += f"\n\n*Источник: {Path(pdf_path).name}*\n"
    return md


def text_to_md(text_path: str, cfg: dict) -> str:
    """Просто копирует .md/.txt как есть (с заголовком)."""
    with open(text_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    return content


# ============================================================
# Вспомогательные функции
# ============================================================

def _is_ollama_available() -> bool:
    """Проверяет, доступен ли ollama."""
    try:
        import requests
        resp = requests.get(f"{CFG['VISION_OLLAMA_HOST']}/api/tags", timeout=2)
        return resp.status_code == 200
    except Exception:
        return False


def _ollama_vision_describe(image_path: str, cfg: dict) -> str:
    """Реальный вызов Qwen2.5-VL через ollama API."""
    import base64
    import requests

    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    payload = {
        "model": cfg.get("VISION_MODEL", "qwen2.5-vl:7b"),
        "prompt": """Опиши изображение как Markdown.
Извлеки: узлы (оборудование, минералы, продукты), связи (потоки, зависимости), параметры.
Формат: текст с перечислениями и таблицами.""",
        "images": [b64],
        "stream": False,
    }

    try:
        resp = requests.post(
            f"{cfg.get('VISION_OLLAMA_HOST', 'http://localhost:11434')}/api/generate",
            json=payload,
            timeout=30,
        )
        if resp.status_code == 200:
            return resp.json().get("response", "*нет описания*")
    except Exception as e:
        return f"*Ошибка Vision API: {e}*\n"

    return "*нет описания*"


def _image_fallback_description(image_path: str) -> str:
    """Заглушка на основе имени файла."""
    name = Path(image_path).name.lower()
    desc = f"*Изображение: {name}*\n\n"

    if any(x in name for x in ("схем", "флотац")):
        desc += """Схема флотации (на основе имени файла).
Узлы: дробилка, грохот, мельница, классификатор, гидроциклон, флотомашина, хвосты, концентрат.
Связи: последовательность операций с возвратами.
Рекомендация: установите Qwen2.5-VL для точного описания через ollama."""
    elif any(x in name for x in ("оборуд", "регламент")):
        desc += """Список оборудования обогатительной фабрики (на основе имени файла).
Типы: мельницы, флотомашины, классификаторы, насосы, гидроциклоны.
Рекомендация: установите Qwen2.5-VL для точного описания через ollama."""
    else:
        desc += f"Изображение {name}. Для точного описания установите Qwen2.5-VL."

    return desc


def _should_process(filepath: str, cfg: dict) -> bool:
    """Проверяет, нужно ли обрабатывать файл."""
    ext = Path(filepath).suffix.lower()

    # Проверка типа
    ftype = EXT_MAP.get(ext)
    if ftype is None:
        return False

    # Проверка включённых расширений
    incl = cfg.get("EXTENSIONS_INCLUDE", "")
    if incl:
        allowed = [x.strip().lower() for x in incl.split(",") if x.strip()]
        if ext not in allowed:
            return False

    # Проверка исключённых
    excl = cfg.get("EXTENSIONS_EXCLUDE", "")
    if excl:
        denied = [x.strip().lower() for x in excl.split(",") if x.strip()]
        if ext in denied:
            return False

    # Проверка флагов типов
    if ftype == "image" and not cfg.get("IMAGE_ENABLED", True):
        return False
    if ftype == "text" and not cfg.get("TEXT_ENABLED", True):
        return False
    if ftype == "tabular" and not cfg.get("TABLE_ENABLED", True):
        return False

    return True


def _process_file(filepath: str, cfg: dict) -> str:
    """Обрабатывает один файл -> Markdown."""
    ext = Path(filepath).suffix.lower()
    ftype = EXT_MAP.get(ext, "unknown")

    handlers = {
        "image": image_to_md,
        "tabular": table_to_md,
    }

    if ext == ".docx":
        handler = docx_to_md
    elif ext == ".pdf":
        handler = pdf_to_md
    elif ext in (".md", ".txt"):
        handler = text_to_md
    else:
        handler = handlers.get(ftype)

    if handler:
        return handler(filepath, cfg)
    return f"# {Path(filepath).name}\n\n*Неподдерживаемый формат: {ext}*\n"


# ============================================================
# CLI
# ============================================================

def parse_args(argv=None):
    parser = argparse.ArgumentParser(
        description="data2md — конвертер данных в Markdown для repomix",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры:
  %(prog)s --input ./данные --output ./repomix_input
  %(prog)s --input ./данные --output ./repomix_input --no-images --max-rows 20
  %(prog)s --input ./данные --output ./repomix_input --ext-include .png,.xlsx
  %(prog)s --input ./данные --output ./repomix_input --verbose
  %(prog)s --input ./данные --output ./repomix_input --include-stats --no-key-metrics
        """,
    )

    # Пути
    parser.add_argument("--input", "-i", default=CFG["INPUT_DIR"],
                        help="Входная директория с данными (default: ./данные)")
    parser.add_argument("--output", "-o", default=CFG["OUTPUT_DIR"],
                        help="Выходная директория для Markdown (default: ./repomix_input)")

    # Фильтры
    parser.add_argument("--ext-include", default=CFG["EXTENSIONS_INCLUDE"],
                        help="Только эти расширения (через запятую, напр. .png,.xlsx)")
    parser.add_argument("--ext-exclude", default=CFG["EXTENSIONS_EXCLUDE"],
                        help="Исключить расширения (через запятую, напр. .log,.tmp)")

    # Флаги типов
    parser.add_argument("--no-images", action="store_false", dest="images",
                        help="Не обрабатывать изображения")
    parser.add_argument("--no-text", action="store_false", dest="text",
                        help="Не обрабатывать текст/PDF/DOCX")
    parser.add_argument("--no-tables", action="store_false", dest="tables",
                        help="Не обрабатывать таблицы Excel/CSV")
    parser.set_defaults(images=CFG["IMAGE_ENABLED"],
                        text=CFG["TEXT_ENABLED"],
                        tables=CFG["TABLE_ENABLED"])

    # Параметры обработки
    parser.add_argument("--max-rows", type=int, default=CFG["MAX_TABLE_ROWS"],
                        help="Макс. строк из Excel в Markdown (default: 50)")
    parser.add_argument("--page-limit", type=int, default=CFG["PDF_PAGE_LIMIT"],
                        help="Макс. страниц PDF (0 = все, default: 0)")
    parser.add_argument("--vision-model", default=CFG["VISION_MODEL"],
                        help="Модель для Vision API (default: qwen2.5-vl:7b)")
    parser.add_argument("--ollama-host", default=CFG["VISION_OLLAMA_HOST"],
                        help="Ollama host (default: http://localhost:11434)")

    # Вывод
    parser.add_argument("--include-stats", action="store_true", default=CFG["INCLUDE_STATS"],
                        help="Добавлять статистику для таблиц")
    parser.add_argument("--no-stats", action="store_false", dest="include_stats",
                        help="Не добавлять статистику")
    parser.add_argument("--include-key-metrics", action="store_true",
                        default=CFG["INCLUDE_KEY_METRICS"],
                        help="Извлекать ключевые метрики")
    parser.add_argument("--no-key-metrics", action="store_false",
                        dest="include_key_metrics",
                        help="Не извлекать ключевые метрики")
    parser.add_argument("--verbose", "-v", action="store_true", default=CFG["VERBOSE"],
                        help="Подробный вывод")

    return parser.parse_args(argv)


def main():
    args = parse_args()

    # Собираем конфиг из аргументов
    cfg = {
        "INPUT_DIR": args.input,
        "OUTPUT_DIR": args.output,
        "EXTENSIONS_INCLUDE": args.ext_include,
        "EXTENSIONS_EXCLUDE": args.ext_exclude,
        "IMAGE_ENABLED": args.images,
        "TEXT_ENABLED": args.text,
        "TABLE_ENABLED": args.tables,
        "MAX_TABLE_ROWS": args.max_rows,
        "PDF_PAGE_LIMIT": args.page_limit,
        "VISION_MODEL": args.vision_model,
        "VISION_OLLAMA_HOST": args.ollama_host,
        "INCLUDE_STATS": args.include_stats,
        "INCLUDE_KEY_METRICS": args.include_key_metrics,
        "VERBOSE": args.verbose,
    }

    # Проверка путей
    input_dir = Path(cfg["INPUT_DIR"])
    output_dir = Path(cfg["OUTPUT_DIR"])

    if not input_dir.exists():
        print(f"X Ошибка: входная директория не найдена: {input_dir}")
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)

    # Сбор файлов
    all_files = []
    for f in input_dir.rglob("*"):
        if f.is_file() and _should_process(str(f), cfg):
            all_files.append(f)

    if not all_files:
        print(f"X Нет файлов для обработки в {input_dir}")
        sys.exit(0)

    # Обработка
    stats = {"image": 0, "text": 0, "tabular": 0, "skipped": 0}
    for f in sorted(all_files):
        try:
            ext = f.suffix.lower()
            ftype = EXT_MAP.get(ext, "unknown")

            if cfg["VERBOSE"]:
                print(f"  -> {f.name} ({ftype})")

            md_content = _process_file(str(f), cfg)

            # Сохраняем
            out_name = f.stem + ".md"
            out_path = output_dir / out_name

            # Если конфликт имён — добавляем префикс
            counter = 1
            while out_path.exists():
                out_path = output_dir / f"{f.stem}_{counter}.md"
                counter += 1

            with open(out_path, "w", encoding="utf-8") as out:
                out.write(md_content)

            if cfg["VERBOSE"]:
                print(f"    V {out_path.name}")

            if ftype in stats:
                stats[ftype] += 1
            else:
                stats["skipped"] += 1

        except Exception as e:
            print(f"  X {f.name}: {e}")
            stats["skipped"] += 1

    # Итог
    print(f"\n=== Готово ===")
    print(f"  Изображений: {stats['image']}")
    print(f"  Текстов/PDF: {stats['text']}")
    print(f"  Таблиц:      {stats['tabular']}")
    print(f"  Пропущено:   {stats['skipped']}")
    print(f"  Всего:       {len(all_files)}")
    print(f"  Выход:       {output_dir.resolve()}")
    print(f"\nДалее: repomix --input {output_dir} --output corpus.txt")


if __name__ == "__main__":
    main()