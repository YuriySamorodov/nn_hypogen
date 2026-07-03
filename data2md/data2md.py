#!/usr/bin/env python3
"""
data2md - конвертер данных любого типа в Markdown для последующей обработки LLM через repomix или yek.

Назначение:
    Преобразует разнородные данные (изображения, PDF, DOCX, Excel, CSV, текст)
    в единый текстовый формат Markdown, который repomix или yek может упаковать в единый
    корпус для подачи в локальную LLM (MetalGPT-1, Qwen2.5-VL).

Поддерживаемые форматы:
    - Изображения: PNG, JPG, JPEG
      -> Текстовое описание через Qwen2.5-VL (ollama), Deepseek API, VK API,
         Yandex Cloud Vision API, или fallback-описание на основе имени файла.
      -> Результат кешируется по хешу файла, чтобы не платить за повторный вызов API.

    - Текст: PDF, DOCX, MD, TXT
      -> Извлечение текста с сохранением структуры (заголовки, параграфы, таблицы).

    - Таблицы: XLSX, XLS, CSV
      -> Чтение через pandas (в т.ч. все листы Excel), вывод в Markdown-таблицы,
         статистика по числовым колонкам, извлечение ключевых метрик
         (Ni, Cu, Fe, потери, класс крупности и т.д.).

Что изменилось по сравнению с версией 1.x:
    - Логирование через logging вместо print (--verbose управляет уровнем).
    - Конфигурация оформлена как dataclass Config вместо словаря.
    - 4 почти идентичные функции вызова Vision API объединены в одну с retry/backoff.
    - Кеш описаний изображений по SHA-256 файла (не дергаем API повторно).
    - Параллельная обработка файлов (--workers).
    - Устойчивое чтение CSV/TXT в разных кодировках (utf-8 / cp1251 / latin-1).
    - Поддержка всех листов Excel, а не только первого.
    - Имена выходных .md файлов сохраняют относительный путь, чтобы не терять
      структуру и не плодить случайные "_1", "_2" при одинаковых именах файлов.
    - Пропуск файлов больше --max-file-size-mb с предупреждением.
    - --dry-run: показать, что будет обработано, без записи файлов.
    - --summary-json: сохранить машиночитаемый отчёт о прогоне.
    - Исходные CLI-флаги полностью сохранены для обратной совместимости.

Использование:
    python data2md.py --input ./данные --output ./repomix_input
    python data2md.py --input ./данные --output ./repomix_input --no-images --max-rows 20
    python data2md.py --input ./данные --output ./repomix_input --ext-include .png,.xlsx --verbose
    python data2md.py --input ./данные --output ./repomix_input --vision-provider deepseek --vision-deepseek-key YOUR_KEY
    python data2md.py --input ./данные --output ./repomix_input --vision-provider vk --vision-vk-key YOUR_KEY
    python data2md.py --input ./данные --output ./repomix_input --vision-provider yandex --vision-yandex-key YOUR_KEY
    python data2md.py --input ./данные --output ./repomix_input --corpus-builder yek
    python data2md.py --input ./данные --output ./repomix_input --workers 8 --dry-run
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import json
import logging
import os
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from pathlib import Path
from threading import Lock
from typing import Callable, Optional

import pandas as pd
from docx import Document
from pypdf import PdfReader

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv не обязателен

try:
    import requests
except ImportError:
    requests = None  # без него Vision-провайдеры недоступны, но остальное работает

try:
    from tqdm import tqdm  # опционально, для прогресс-бара
except ImportError:
    tqdm = None

VERSION = "2.0.0"

logger = logging.getLogger("data2md")

# ============================================================
# Статичные справочники
# ============================================================

EXT_MAP = {
    ".png": "image", ".jpg": "image", ".jpeg": "image",
    ".pdf": "text",
    ".docx": "text",
    ".md": "text", ".txt": "text",
    ".xlsx": "tabular", ".xls": "tabular", ".csv": "tabular",
}

KEY_METRICS_KEYWORDS = [
    'ni', 'никель', 'cu', 'медь', 'потер', 'содерж',
    'co', 'кобальт', 'fe', 'железо', 's', 'сера',
    'au', 'золото', 'ag', 'серебро', 'pt', 'pd',
    'класс', 'крупность', 'извлечен', 'хвост',
]

# Кодировки, которые пробуем по очереди при чтении текстовых файлов.
TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1251", "latin-1")

VISION_PROMPT = (
    "Опиши изображение как Markdown.\n"
    "Извлеки: узлы (оборудование, минералы, продукты), связи (потоки, зависимости), параметры.\n"
    "Формат: текст с перечислениями и таблицами."
)

# Конфигурация HTTP Vision-провайдеров (кроме ollama — у него свой протокол).
# Каждый провайдер описывает: URL, как собрать заголовки, тело запроса и как
# извлечь текст ответа. Это заменяет 3 дублирующиеся функции из версии 1.x.
HTTP_VISION_PROVIDERS: dict = {
    "deepseek": {
        "url": "https://api.deepseek.com/v1/chat/completions",
        "model": "deepseek-vl",
        "headers": lambda key: {"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
        "payload": lambda b64, model: {
            "model": model,
            "messages": [{"role": "user", "content": VISION_PROMPT, "images": [b64]}],
            "stream": False,
        },
        "parse": lambda data: data.get("choices", [{}])[0].get("message", {}).get("content"),
    },
    "vk": {
        "url": "https://api.vk.com/v1/images/describe",
        "model": "vision",
        "headers": lambda key: {"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
        "payload": lambda b64, model: {
            "model": model,
            "messages": [{"role": "user", "content": VISION_PROMPT, "images": [b64]}],
            "stream": False,
        },
        "parse": lambda data: data.get("description"),
    },
    "yandex": {
        "url": "https://vision.api.cloud.yandex.net/v1/images/describe",
        "model": "vision",
        "headers": lambda key: {"Authorization": f"Api-Key {key}", "Content-Type": "application/json"},
        "payload": lambda b64, model: {
            "model": model,
            "messages": [{"role": "user", "content": VISION_PROMPT, "images": [b64]}],
            "stream": False,
        },
        "parse": lambda data: data.get("description"),
    },
}


# ============================================================
# Конфигурация запуска
# ============================================================

@dataclass
class Config:
    input_dir: Path
    output_dir: Path
    corpus_builder: str = "repomix"

    ext_include: list = field(default_factory=list)
    ext_exclude: list = field(default_factory=lambda: [".log", ".tmp"])

    images_enabled: bool = True
    text_enabled: bool = True
    tables_enabled: bool = True

    max_table_rows: int = 50
    pdf_page_limit: int = 0

    vision_provider: str = "ollama"
    vision_model: str = "qwen2.5-vl:7b"
    ollama_host: str = "http://localhost:11434"
    vision_keys: dict = field(default_factory=dict)  # provider -> api key

    include_stats: bool = True
    include_key_metrics: bool = True

    verbose: bool = False
    workers: int = 4
    max_file_size_mb: float = 200.0
    dry_run: bool = False
    cache_path: Optional[Path] = None
    summary_json: Optional[Path] = None

    def should_process(self, filepath: Path) -> bool:
        ext = filepath.suffix.lower()
        ftype = EXT_MAP.get(ext)
        if ftype is None:
            return False
        if self.ext_include and ext not in self.ext_include:
            return False
        if ext in self.ext_exclude:
            return False
        if ftype == "image" and not self.images_enabled:
            return False
        if ftype == "text" and not self.text_enabled:
            return False
        if ftype == "tabular" and not self.tables_enabled:
            return False
        return True


# ============================================================
# Кеш описаний изображений (по SHA-256 содержимого файла)
# ============================================================

class VisionCache:
    """Простой JSON-кеш на диске, чтобы не пересчитывать описания изображений
    при повторных запусках (Vision API может быть платным/медленным)."""

    def __init__(self, path: Optional[Path]):
        self.path = path
        self._lock = Lock()
        self._data: dict = {}
        if self.path and self.path.exists():
            try:
                self._data = json.loads(self.path.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError) as e:
                logger.warning(f"Не удалось прочитать кеш {self.path}: {e}")
                self._data = {}

    @staticmethod
    def _key(file_hash: str, provider: str, model: str) -> str:
        return f"{provider}:{model}:{file_hash}"

    def get(self, file_hash: str, provider: str, model: str) -> Optional[str]:
        return self._data.get(self._key(file_hash, provider, model))

    def set(self, file_hash: str, provider: str, model: str, value: str) -> None:
        with self._lock:
            self._data[self._key(file_hash, provider, model)] = value
            if self.path:
                try:
                    self.path.write_text(
                        json.dumps(self._data, ensure_ascii=False, indent=2), encoding="utf-8"
                    )
                except OSError as e:
                    logger.warning(f"Не удалось сохранить кеш {self.path}: {e}")


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


# ============================================================
# Обработчики форматов
# ============================================================

def image_to_md(image_path: Path, cfg: Config, cache: VisionCache) -> str:
    """Изображение -> Markdown-описание через Vision API (с кешированием)."""
    name = image_path.name
    md = f"# {name}\n\n"

    provider = cfg.vision_provider
    try:
        file_hash = _sha256_file(image_path)
    except OSError as e:
        return md + f"*Не удалось прочитать файл: {e}*\n"

    cached = cache.get(file_hash, provider, cfg.vision_model)
    if cached is not None:
        logger.debug(f"{name}: описание взято из кеша")
        md += cached
        md += f"\n\n*Источник: {name} (кеш)*\n"
        return md

    if requests is None:
        description = _image_fallback_description(image_path)
    elif provider == "ollama" and _is_ollama_available(cfg.ollama_host):
        description = _ollama_vision_describe(image_path, cfg)
    elif provider in HTTP_VISION_PROVIDERS and cfg.vision_keys.get(provider):
        description = _http_vision_describe(image_path, cfg, provider)
    else:
        description = _image_fallback_description(image_path)

    md += description
    md += f"\n\n*Источник: {name}*\n"

    # Кешируем только настоящие ответы API, не fallback-заглушку и не ошибки.
    if requests is not None and not description.startswith("*Изображение") and not description.startswith("*Ошибка"):
        cache.set(file_hash, provider, cfg.vision_model, description)

    return md


def _read_excel_all_sheets(path: Path) -> dict:
    """Читает все листы Excel-файла. Возвращает {имя_листа: DataFrame}."""
    try:
        return pd.read_excel(path, sheet_name=None)
    except ImportError as e:
        # Например, отсутствует xlrd для старого .xls
        raise RuntimeError(
            f"Не хватает библиотеки для чтения {path.suffix}: {e}. "
            f"Для .xls установите: pip install xlrd"
        ) from e


def _read_csv_resilient(path: Path) -> pd.DataFrame:
    """Пробует разные кодировки, т.к. русские CSV часто в cp1251."""
    last_exc = None
    for enc in TEXT_ENCODINGS:
        try:
            return pd.read_csv(path, encoding=enc)
        except (UnicodeDecodeError, UnicodeError) as e:
            last_exc = e
            continue
    # Последняя попытка с заменой нечитаемых символов, чтобы не падать совсем.
    logger.warning(f"{path.name}: не удалось однозначно определить кодировку, читаю с заменой символов")
    return pd.read_csv(path, encoding="utf-8", encoding_errors="replace")


def table_to_md(table_path: Path, cfg: Config) -> str:
    """Excel (все листы) / CSV -> Markdown таблицы + статистика + ключевые метрики."""
    ext = table_path.suffix.lower()
    md = f"# {table_path.name}\n\n"

    try:
        if ext == ".csv":
            sheets = {"": _read_csv_resilient(table_path)}
        else:
            sheets = _read_excel_all_sheets(table_path)
    except Exception as e:
        logger.error(f"{table_path.name}: ошибка загрузки — {e}")
        return md + f"*Ошибка загрузки: {e}*\n"

    for sheet_name, df in sheets.items():
        if sheet_name:
            md += f"## Лист: {sheet_name}\n\n"

        md += f"- **Строк**: {len(df)}, **Колонок**: {len(df.columns)}\n"
        dtypes_preview = ', '.join(str(df[c].dtype) for c in df.columns[:5])
        md += f"- **Типы данных**: {dtypes_preview}"
        if len(df.columns) > 5:
            md += f" ... (+{len(df.columns) - 5})"
        md += "\n\n"

        md += "### Колонки\n\n"
        for i, col in enumerate(df.columns):
            non_null = df[col].notna().sum()
            md += f"{i + 1}. `{col}` ({df[col].dtype}) — {non_null}/{len(df)} непустых\n"

        max_rows = cfg.max_table_rows
        md += f"\n### Данные (первые {min(max_rows, len(df))} строк)\n\n"
        try:
            md += df.head(max_rows).to_markdown(index=False)
        except ImportError:
            # tabulate не установлен — fallback на plain-text представление
            md += df.head(max_rows).to_string(index=False)

        if cfg.include_stats:
            numeric_cols = df.select_dtypes(include="number").columns
            if len(numeric_cols) > 0:
                md += "\n\n### Статистика\n\n"
                stats = df[numeric_cols].describe().round(3)
                stats.index.name = "метрика"
                try:
                    md += stats.to_markdown()
                except ImportError:
                    md += stats.to_string()

        if cfg.include_key_metrics:
            for col in df.columns:
                col_lower = str(col).lower()
                if any(kw in col_lower for kw in KEY_METRICS_KEYWORDS):
                    num = pd.to_numeric(df[col], errors='coerce')
                    valid = num.dropna()
                    if len(valid) > 0:
                        md += f"\n\n**Ключевая колонка**: `{col}`\n"
                        md += f"- Мин: {valid.min():.3f}\n"
                        md += f"- Макс: {valid.max():.3f}\n"
                        md += f"- Среднее: {valid.mean():.3f}\n"
                        md += f"- Медиана: {valid.median():.3f}\n"
                        md += f"- Сумма: {valid.sum():.3f}\n"

        md += "\n\n"

    md += f"*Источник: {table_path.name}*\n"
    return md


def docx_to_md(docx_path: Path, cfg: Config) -> str:
    """DOCX -> Markdown."""
    doc = Document(docx_path)
    md = f"# {docx_path.name}\n\n"
    md += f"**Параграфов**: {len(doc.paragraphs)}, **Таблиц**: {len(doc.tables)}\n\n"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower() if para.style else ""
        if "heading" in style or "заголов" in style:
            digits = "".join(c for c in style if c.isdigit())
            level = min(int(digits) + 1, 6) if digits else 2
            md += f"{'#' * level} {text}\n\n"
        else:
            md += text + "\n\n"

    for i, table in enumerate(doc.tables):
        md += f"## Таблица {i + 1}\n\n"
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            header, data = rows[0], rows[1:]
            if data:
                try:
                    tdf = pd.DataFrame(data, columns=header)
                    md += tdf.to_markdown(index=False)
                except (ImportError, ValueError):
                    md += f"| {' | '.join(header)} |\n"
                    for row in data:
                        md += f"| {' | '.join(row)} |\n"
            else:
                md += f"| {' | '.join(header)} |\n"
        md += "\n"

    md += f"\n\n*Источник: {docx_path.name}*\n"
    return md


def pdf_to_md(pdf_path: Path, cfg: Config) -> str:
    """PDF -> Markdown."""
    try:
        reader = PdfReader(pdf_path)
    except Exception as e:
        return f"# {pdf_path.name}\n\n*Ошибка чтения PDF: {e}*\n"

    total = len(reader.pages)
    limit = cfg.pdf_page_limit
    pages = total if limit <= 0 else min(total, limit)

    md = f"# {pdf_path.name}\n\n"
    md += f"**Страниц**: {total}, обработано: {pages}\n\n"

    for i in range(pages):
        try:
            text = reader.pages[i].extract_text() or ""
        except Exception as e:
            logger.warning(f"{pdf_path.name}: страница {i + 1} не читается — {e}")
            continue
        if text.strip():
            md += f"## Страница {i + 1}\n\n{text}\n\n"

    md += f"\n\n*Источник: {pdf_path.name}*\n"
    return md


def text_to_md(text_path: Path, cfg: Config) -> str:
    """Копирует .md/.txt как есть, пробуя разные кодировки."""
    for enc in TEXT_ENCODINGS:
        try:
            return text_path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return text_path.read_text(encoding="utf-8", errors="replace")


# ============================================================
# Vision API — вспомогательные функции
# ============================================================

def _is_ollama_available(host: str) -> bool:
    if requests is None:
        return False
    try:
        resp = requests.get(f"{host}/api/tags", timeout=2)
        return resp.status_code == 200
    except requests.RequestException:
        return False


def _post_with_retry(url: str, headers: dict, payload: dict, timeout: int = 30,
                      retries: int = 2, backoff: float = 1.5):
    """POST с несколькими попытками и экспоненциальной задержкой."""
    last_exc: Optional[Exception] = None
    for attempt in range(retries + 1):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=timeout)
            resp.raise_for_status()
            return resp
        except requests.RequestException as e:
            last_exc = e
            if attempt < retries:
                wait = backoff ** attempt
                logger.debug(f"Попытка {attempt + 1} неудачна ({e}), повтор через {wait:.1f}с")
                time.sleep(wait)
    raise last_exc  # type: ignore[misc]


def _ollama_vision_describe(image_path: Path, cfg: Config) -> str:
    """Вызов Qwen2.5-VL через локальный ollama."""
    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    payload = {
        "model": cfg.vision_model,
        "prompt": VISION_PROMPT,
        "images": [b64],
        "stream": False,
    }

    try:
        resp = _post_with_retry(f"{cfg.ollama_host}/api/generate", headers={}, payload=payload)
        return resp.json().get("response", "*нет описания*")
    except Exception as e:
        logger.warning(f"{image_path.name}: ошибка Ollama Vision — {e}")
        return f"*Ошибка Vision API (ollama): {e}*\n"


def _http_vision_describe(image_path: Path, cfg: Config, provider: str) -> str:
    """Единая функция для Deepseek / VK / Yandex — провайдеры отличаются
    только URL/заголовками/форматом ответа, описанными в HTTP_VISION_PROVIDERS."""
    spec = HTTP_VISION_PROVIDERS[provider]
    api_key = cfg.vision_keys.get(provider, "")

    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    headers = spec["headers"](api_key)
    payload = spec["payload"](b64, cfg.vision_model or spec["model"])

    try:
        resp = _post_with_retry(spec["url"], headers=headers, payload=payload)
        data = resp.json()
        result = spec["parse"](data)
        return result or "*нет описания*"
    except Exception as e:
        logger.warning(f"{image_path.name}: ошибка {provider} Vision — {e}")
        return f"*Ошибка {provider} API: {e}*\n"


def _image_fallback_description(image_path: Path) -> str:
    """Заглушка на основе имени файла, когда Vision API недоступен."""
    name = image_path.name.lower()
    desc = f"*Изображение: {name}*\n\n"

    if any(x in name for x in ("схем", "флотац")):
        desc += (
            "Схема флотации (на основе имени файла).\n"
            "Узлы: дробилка, грохот, мельница, классификатор, гидроциклон, флотомашина, хвосты, концентрат.\n"
            "Связи: последовательность операций с возвратами.\n"
            "Рекомендация: установите Qwen2.5-VL для точного описания через ollama."
        )
    elif any(x in name for x in ("оборуд", "регламент")):
        desc += (
            "Список оборудования обогатительной фабрики (на основе имени файла).\n"
            "Типы: мельницы, флотомашины, классификаторы, насосы, гидроциклоны.\n"
            "Рекомендация: установите Qwen2.5-VL для точного описания через ollama."
        )
    else:
        desc += f"Изображение {name}. Для точного описания установите Qwen2.5-VL."

    return desc


# ============================================================
# Обработка файлов
# ============================================================

HANDLERS: dict = {
    "tabular": lambda path, cfg, cache: table_to_md(path, cfg),
    "image": lambda path, cfg, cache: image_to_md(path, cfg, cache),
}


def _process_file(filepath: Path, cfg: Config, cache: VisionCache) -> str:
    ext = filepath.suffix.lower()

    if ext == ".docx":
        return docx_to_md(filepath, cfg)
    if ext == ".pdf":
        return pdf_to_md(filepath, cfg)
    if ext in (".md", ".txt"):
        return text_to_md(filepath, cfg)

    ftype = EXT_MAP.get(ext, "unknown")
    handler = HANDLERS.get(ftype)
    if handler:
        return handler(filepath, cfg, cache)
    return f"# {filepath.name}\n\n*Неподдерживаемый формат: {ext}*\n"


def _output_path_for(filepath: Path, input_dir: Path, output_dir: Path) -> Path:
    """Строит имя выходного файла, сохраняя относительный путь (заменяя
    разделители на '__'), чтобы файлы с одинаковым именем в разных
    подпапках не затирали друг друга и не превращались в бессмысленный _1/_2."""
    try:
        rel = filepath.relative_to(input_dir)
    except ValueError:
        rel = Path(filepath.name)
    flat_name = "__".join(rel.with_suffix("").parts)
    out_path = output_dir / f"{flat_name}.md"

    counter = 1
    while out_path.exists():
        out_path = output_dir / f"{flat_name}_{counter}.md"
        counter += 1
    return out_path


@dataclass
class FileResult:
    path: Path
    ftype: str
    ok: bool
    out_path: Optional[Path] = None
    error: Optional[str] = None
    skipped_reason: Optional[str] = None


def _handle_one(filepath: Path, cfg: Config, cache: VisionCache) -> FileResult:
    ext = filepath.suffix.lower()
    ftype = EXT_MAP.get(ext, "unknown")

    size_mb = filepath.stat().st_size / (1024 * 1024)
    if size_mb > cfg.max_file_size_mb:
        msg = f"пропущен: {size_mb:.1f} МБ > лимита {cfg.max_file_size_mb} МБ"
        logger.warning(f"{filepath.name}: {msg}")
        return FileResult(filepath, ftype, ok=False, skipped_reason=msg)

    if cfg.dry_run:
        logger.info(f"[dry-run] {filepath} ({ftype})")
        return FileResult(filepath, ftype, ok=True)

    try:
        md_content = _process_file(filepath, cfg, cache)
        out_path = _output_path_for(filepath, cfg.input_dir, cfg.output_dir)
        out_path.write_text(md_content, encoding="utf-8")
        logger.debug(f"{filepath.name} -> {out_path.name}")
        return FileResult(filepath, ftype, ok=True, out_path=out_path)
    except Exception as e:
        logger.error(f"{filepath.name}: {e}")
        return FileResult(filepath, ftype, ok=False, error=str(e))


# ============================================================
# CLI
# ============================================================

def parse_args(argv=None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="data2md — конвертер данных в Markdown для repomix",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры:
  %(prog)s --input ./данные --output ./repomix_input
  %(prog)s --input ./данные --output ./repomix_input --no-images --max-rows 20
  %(prog)s --input ./данные --output ./repomix_input --ext-include .png,.xlsx
  %(prog)s --input ./данные --output ./repomix_input --verbose
  %(prog)s --input ./данные --output ./repomix_input --workers 8 --dry-run
        """,
    )

    parser.add_argument("--input", "-i", default=os.getenv("INPUT_DIR", "./данные"),
                         help="Входная директория с данными (default: ./данные)")
    parser.add_argument("--output", "-o", default=os.getenv("OUTPUT_DIR", "./repomix_input"),
                         help="Выходная директория для Markdown (default: ./repomix_input)")

    parser.add_argument("--ext-include", default=os.getenv("EXTENSIONS_INCLUDE", ""),
                         help="Только эти расширения (через запятую, напр. .png,.xlsx)")
    parser.add_argument("--ext-exclude", default=os.getenv("EXTENSIONS_EXCLUDE", ".log,.tmp"),
                         help="Исключить расширения (через запятую, напр. .log,.tmp)")

    parser.add_argument("--no-images", action="store_false", dest="images",
                         help="Не обрабатывать изображения")
    parser.add_argument("--no-text", action="store_false", dest="text",
                         help="Не обрабатывать текст/PDF/DOCX")
    parser.add_argument("--no-tables", action="store_false", dest="tables",
                         help="Не обрабатывать таблицы Excel/CSV")
    parser.set_defaults(
        images=os.getenv("IMAGE_ENABLED", "true").lower() == "true",
        text=os.getenv("TEXT_ENABLED", "true").lower() == "true",
        tables=os.getenv("TABLE_ENABLED", "true").lower() == "true",
    )

    parser.add_argument("--max-rows", type=int, default=int(os.getenv("MAX_TABLE_ROWS", "50")),
                         help="Макс. строк из Excel в Markdown (default: 50)")
    parser.add_argument("--page-limit", type=int, default=int(os.getenv("PDF_PAGE_LIMIT", "0")),
                         help="Макс. страниц PDF (0 = все, default: 0)")
    parser.add_argument("--vision-model", default=os.getenv("VISION_MODEL", "qwen2.5-vl:7b"),
                         help="Модель для Vision API (default: qwen2.5-vl:7b)")
    parser.add_argument("--ollama-host", default=os.getenv("VISION_OLLAMA_HOST", "http://localhost:11434"),
                         help="Ollama host (default: http://localhost:11434)")
    parser.add_argument("--vision-provider", default=os.getenv("VISION_PROVIDER", "ollama"),
                         choices=["ollama", "deepseek", "vk", "yandex"],
                         help="Провайдер Vision API (default: ollama)")
    parser.add_argument("--corpus-builder", default=os.getenv("CORPUS_BUILDER", "repomix"),
                         choices=["repomix", "yek"],
                         help="Корпусный строитель (default: repomix)")
    parser.add_argument("--vision-deepseek-key", default=os.getenv("VISION_DEEPSEEK_API_KEY", ""),
                         help="API ключ для Deepseek Vision")
    parser.add_argument("--vision-vk-key", default=os.getenv("VISION_VK_API_KEY", ""),
                         help="API ключ для VK Vision")
    parser.add_argument("--vision-yandex-key", default=os.getenv("VISION_YANDEX_CLOUD_API_KEY", ""),
                         help="API ключ для Yandex Cloud Vision")

    parser.add_argument("--include-stats", action="store_true",
                         default=os.getenv("INCLUDE_STATS", "true").lower() == "true",
                         help="Добавлять статистику для таблиц")
    parser.add_argument("--no-stats", action="store_false", dest="include_stats",
                         help="Не добавлять статистику")
    parser.add_argument("--include-key-metrics", action="store_true",
                         default=os.getenv("INCLUDE_KEY_METRICS", "true").lower() == "true",
                         help="Извлекать ключевые метрики")
    parser.add_argument("--no-key-metrics", action="store_false", dest="include_key_metrics",
                         help="Не извлекать ключевые метрики")
    parser.add_argument("--verbose", "-v", action="store_true",
                         default=os.getenv("VERBOSE", "false").lower() == "true",
                         help="Подробный вывод")

    # Новые параметры
    parser.add_argument("--workers", type=int, default=int(os.getenv("WORKERS", "4")),
                         help="Число потоков для параллельной обработки (default: 4)")
    parser.add_argument("--max-file-size-mb", type=float, default=float(os.getenv("MAX_FILE_SIZE_MB", "200")),
                         help="Пропускать файлы больше указанного размера в МБ (default: 200)")
    parser.add_argument("--dry-run", action="store_true",
                         help="Показать, что будет обработано, без записи файлов")
    parser.add_argument("--cache-file", default=os.getenv("VISION_CACHE_FILE", ""),
                         help="Путь к файлу кеша описаний изображений "
                              "(default: <output>/.data2md_vision_cache.json)")
    parser.add_argument("--no-cache", action="store_true",
                         help="Отключить кеширование описаний изображений")
    parser.add_argument("--summary-json", default="",
                         help="Сохранить машиночитаемый отчёт о прогоне в JSON")
    parser.add_argument("--version", action="version", version=f"data2md {VERSION}")

    return parser.parse_args(argv)


def build_config(args: argparse.Namespace) -> Config:
    return Config(
        input_dir=Path(args.input),
        output_dir=Path(args.output),
        corpus_builder=args.corpus_builder,
        ext_include=[x.strip().lower() for x in args.ext_include.split(",") if x.strip()],
        ext_exclude=[x.strip().lower() for x in args.ext_exclude.split(",") if x.strip()],
        images_enabled=args.images,
        text_enabled=args.text,
        tables_enabled=args.tables,
        max_table_rows=args.max_rows,
        pdf_page_limit=args.page_limit,
        vision_provider=args.vision_provider,
        vision_model=args.vision_model,
        ollama_host=args.ollama_host,
        vision_keys={
            "deepseek": args.vision_deepseek_key,
            "vk": args.vision_vk_key,
            "yandex": args.vision_yandex_key,
        },
        include_stats=args.include_stats,
        include_key_metrics=args.include_key_metrics,
        verbose=args.verbose,
        workers=max(1, args.workers),
        max_file_size_mb=args.max_file_size_mb,
        dry_run=args.dry_run,
        cache_path=None if args.no_cache else Path(
            args.cache_file or (Path(args.output) / ".data2md_vision_cache.json")
        ),
        summary_json=Path(args.summary_json) if args.summary_json else None,
    )


def setup_logging(verbose: bool) -> None:
    logging.basicConfig(
        level=logging.DEBUG if verbose else logging.INFO,
        format="%(message)s",
        stream=sys.stdout,
    )


def main(argv=None) -> int:
    args = parse_args(argv)
    setup_logging(args.verbose)
    cfg = build_config(args)

    if not cfg.input_dir.exists():
        logger.error(f"X Ошибка: входная директория не найдена: {cfg.input_dir}")
        return 1

    try:
        cfg.input_dir.resolve().relative_to(cfg.output_dir.resolve())
        logger.error("X Ошибка: выходная директория не должна быть родительской для входной")
        return 1
    except ValueError:
        pass  # это нормальный случай — input не внутри output

    if not cfg.dry_run:
        cfg.output_dir.mkdir(parents=True, exist_ok=True)

    all_files = sorted(f for f in cfg.input_dir.rglob("*") if f.is_file() and cfg.should_process(f))

    if not all_files:
        logger.info(f"X Нет файлов для обработки в {cfg.input_dir}")
        return 0

    if cfg.vision_provider != "ollama" and cfg.vision_provider in HTTP_VISION_PROVIDERS \
            and not cfg.vision_keys.get(cfg.vision_provider):
        logger.warning(
            f"! Провайдер '{cfg.vision_provider}' выбран, но ключ API не задан — "
            f"изображения получат только fallback-описание по имени файла."
        )

    cache = VisionCache(cfg.cache_path)

    logger.info(f"Найдено файлов: {len(all_files)}. Потоков: {cfg.workers}.")

    results: list[FileResult] = []
    progress = tqdm(total=len(all_files), disable=tqdm is None or cfg.verbose, unit="файл")

    with ThreadPoolExecutor(max_workers=cfg.workers) as pool:
        futures = {pool.submit(_handle_one, f, cfg, cache): f for f in all_files}
        for future in as_completed(futures):
            result = future.result()
            results.append(result)
            if cfg.verbose:
                status = "V" if result.ok else "X"
                extra = result.skipped_reason or result.error or ""
                logger.debug(f"  {status} {result.path.name} {extra}")
            if tqdm is not None:
                progress.update(1)
    if tqdm is not None:
        progress.close()

    stats = {"image": 0, "text": 0, "tabular": 0, "unknown": 0, "skipped": 0, "errors": 0}
    for r in results:
        if r.skipped_reason:
            stats["skipped"] += 1
        elif not r.ok:
            stats["errors"] += 1
        else:
            stats[r.ftype] = stats.get(r.ftype, 0) + 1

    logger.info("\n=== Готово ===")
    logger.info(f"  Изображений: {stats['image']}")
    logger.info(f"  Текстов/PDF: {stats['text']}")
    logger.info(f"  Таблиц:      {stats['tabular']}")
    logger.info(f"  Пропущено:   {stats['skipped']}")
    logger.info(f"  Ошибок:      {stats['errors']}")
    logger.info(f"  Всего:       {len(all_files)}")
    if not cfg.dry_run:
        logger.info(f"  Результат:   {cfg.output_dir.resolve()}")

    if cfg.summary_json:
        summary = {
            "version": VERSION,
            "input_dir": str(cfg.input_dir.resolve()),
            "output_dir": str(cfg.output_dir.resolve()),
            "dry_run": cfg.dry_run,
            "stats": stats,
            "files": [
                {
                    "path": str(r.path),
                    "type": r.ftype,
                    "ok": r.ok,
                    "output": str(r.out_path) if r.out_path else None,
                    "error": r.error,
                    "skipped_reason": r.skipped_reason,
                }
                for r in results
            ],
        }
        cfg.summary_json.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
        logger.info(f"  Отчёт:       {cfg.summary_json.resolve()}")

    if cfg.dry_run:
        logger.info("\n(dry-run — файлы не записывались)")
    else:
        builder = cfg.corpus_builder
        logger.info(f"\nДалее: {builder} --input {cfg.output_dir} --output corpus.txt")

    return 1 if stats["errors"] and stats["errors"] == len(all_files) else 0


if __name__ == "__main__":
    sys.exit(main())
