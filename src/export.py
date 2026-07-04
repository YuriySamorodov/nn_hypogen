from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.schemas import HypothesisBatch


def batch_to_markdown_bytes(batch: HypothesisBatch) -> bytes:
    from src.hypotheses import hypotheses_to_markdown

    return hypotheses_to_markdown(batch).encode("utf-8")


def batch_to_docx_bytes(batch: HypothesisBatch) -> bytes:
    document = Document()
    title = document.add_heading(f"Гипотезы: {batch.target_property}", level=0)
    title.runs[0].font.size = Pt(16)
    document.add_paragraph(f"Ограничения: {batch.constraints}")
    document.add_paragraph(f"Дата: {datetime.now():%Y-%m-%d %H:%M}")

    for index, hypothesis in enumerate(batch.hypotheses, start=1):
        document.add_heading(f"{index}. {hypothesis.title}", level=1)
        document.add_paragraph(f"Формулировка: {hypothesis.statement}")
        document.add_paragraph(f"Обоснование: {hypothesis.rationale}")
        document.add_paragraph(f"Механизм: {hypothesis.mechanism}")
        document.add_paragraph(
            f"Оценки: новизна {hypothesis.novelty_score}/5, "
            f"реализуемость {hypothesis.feasibility_score}/5, "
            f"риск {hypothesis.risk_level}"
        )
        document.add_paragraph(f"Ожидаемый KPI: {hypothesis.expected_kpi_impact}")

        document.add_paragraph("Источники:")
        for evidence in hypothesis.evidence:
            marker = "" if evidence.verified else " [источник не подтверждён]"
            document.add_paragraph(
                f"• {evidence.source_id} — {evidence.title}: "
                f"{evidence.quote} ({evidence.relevance}){marker}",
                style="List Bullet",
            )

        document.add_paragraph("План проверки:")
        for step in hypothesis.validation_plan:
            document.add_paragraph(step, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def write_export_files(batch: HypothesisBatch, directory: Path) -> tuple[Path, Path]:
    directory.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    md_path = directory / f"hypotheses_{stamp}.md"
    docx_path = directory / f"hypotheses_{stamp}.docx"
    md_path.write_bytes(batch_to_markdown_bytes(batch))
    docx_path.write_bytes(batch_to_docx_bytes(batch))
    return md_path, docx_path
