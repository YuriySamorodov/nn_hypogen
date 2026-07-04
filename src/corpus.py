from dataclasses import dataclass
from pathlib import Path
from zipfile import BadZipFile

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader


DEFAULT_CORPUS_DIR = Path("data/demo_corpus")
DEFAULT_TASK_DIR = Path("Задача 1")
LFS_POINTER_PREFIX = "version https://git-lfs.github.com/spec/v1"


def is_lfs_pointer(path: Path) -> bool:
    """True if the file is a Git LFS pointer instead of real content."""
    if not path.is_file():
        return False
    if path.stat().st_size > 512:
        return False
    try:
        header = path.read_text(encoding="utf-8", errors="ignore")[:64]
    except OSError:
        return False
    return header.startswith(LFS_POINTER_PREFIX)


@dataclass(frozen=True)
class CorpusDocument:
    text: str
    metadata: dict[str, str]
    path: Path


REQUIRED_METADATA = {
    "source_id",
    "title",
    "language",
    "domain",
    "material",
    "process",
    "page_or_section",
}


def _parse_frontmatter(raw: str, path: Path) -> tuple[dict[str, str], str]:
    if not raw.startswith("---\n"):
        raise ValueError(f"{path} must start with frontmatter")

    try:
        _, frontmatter, body = raw.split("---", 2)
    except ValueError as exc:
        raise ValueError(f"{path} has invalid frontmatter") from exc

    metadata: dict[str, str] = {}
    for line in frontmatter.strip().splitlines():
        if not line.strip():
            continue
        key, sep, value = line.partition(":")
        if not sep:
            raise ValueError(f"{path} has invalid metadata line: {line}")
        metadata[key.strip()] = value.strip()

    missing = REQUIRED_METADATA - metadata.keys()
    if missing:
        missing_list = ", ".join(sorted(missing))
        raise ValueError(f"{path} missing metadata: {missing_list}")

    return metadata, body.strip()


def load_demo_corpus(corpus_dir: Path = DEFAULT_CORPUS_DIR) -> list[CorpusDocument]:
    documents: list[CorpusDocument] = []
    for path in sorted(corpus_dir.glob("*.md")):
        metadata, body = _parse_frontmatter(path.read_text(encoding="utf-8"), path)
        documents.append(CorpusDocument(text=body, metadata=metadata, path=path))

    if not documents:
        raise ValueError(f"No markdown documents found in {corpus_dir}")

    return documents


def _metadata_for_path(path: Path, source_id: str, title: str, source_type: str) -> dict[str, str]:
    return {
        "source_id": source_id,
        "title": title,
        "language": "ru",
        "domain": "mineral_processing",
        "material": "tailings",
        "process": "flotation",
        "page_or_section": source_type,
    }


def _short_id(prefix: str, path: Path, suffix: str | int | None = None) -> str:
    safe = "".join(ch.upper() if ch.isalnum() else "-" for ch in path.stem)
    safe = "-".join(part for part in safe.split("-") if part)[:36]
    if suffix is None:
        return f"{prefix}-{safe}"
    return f"{prefix}-{safe}-{suffix}"


def _load_docx(path: Path) -> list[CorpusDocument]:
    doc = DocxDocument(path)
    parts: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        return []

    return [
        CorpusDocument(
            text=text,
            metadata=_metadata_for_path(path, _short_id("DOCX", path), path.stem, "docx"),
            path=path,
        )
    ]


def _load_pdf(path: Path) -> list[CorpusDocument]:
    documents: list[CorpusDocument] = []
    reader = PdfReader(str(path))
    for page_index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if len(text) < 80:
            continue
        documents.append(
            CorpusDocument(
                text=text,
                metadata=_metadata_for_path(
                    path,
                    _short_id("PDF", path, page_index),
                    path.stem,
                    f"pdf page {page_index}",
                ),
                path=path,
            )
        )
    return documents


def _row_text(values: list[object]) -> str:
    return " | ".join(str(value).strip() for value in values if value is not None and str(value).strip())


def _load_xlsx(path: Path) -> list[CorpusDocument]:
    documents: list[CorpusDocument] = []
    workbook = load_workbook(path, read_only=True, data_only=True)
    for worksheet in workbook.worksheets:
        rows: list[str] = []
        for row_index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
            text = _row_text(list(row))
            if text:
                rows.append(f"row {row_index}: {text}")

        if rows:
            title = f"{path.parent.name}: {path.stem} / {worksheet.title}"
            documents.append(
                CorpusDocument(
                    text="\n".join(rows),
                    metadata=_metadata_for_path(
                        path,
                        _short_id("XLSX", path, worksheet.title),
                        title,
                        f"xlsx sheet {worksheet.title}",
                    ),
                    path=path,
                )
            )
    return documents


def _try_ocr_png(path: Path, *, languages: str = "rus+eng") -> str | None:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return None

    try:
        text = pytesseract.image_to_string(Image.open(path), lang=languages)
    except (OSError, RuntimeError, ValueError):
        return None

    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    return text if len(text) >= 20 else None


def _load_png(path: Path, *, enable_ocr: bool = True, ocr_languages: str = "rus+eng") -> list[CorpusDocument]:
    title = path.stem
    ocr_text = _try_ocr_png(path, languages=ocr_languages) if enable_ocr else None
    if ocr_text:
        text = (
            f"Визуальный источник: {title}. Папка: {path.parent.name}.\n"
            f"Распознанный текст (OCR):\n{ocr_text}"
        )
        section = "image ocr"
    else:
        text = (
            f"Визуальный источник: {title}. Файл расположен в папке {path.parent.name}. "
            "Это схема, регламент или список оборудования. "
            "OCR недоступен или не распознал достаточно текста."
        )
        section = "image metadata"
    return [
        CorpusDocument(
            text=text,
            metadata=_metadata_for_path(path, _short_id("PNG", path), title, section),
            path=path,
        )
    ]


def load_task_corpus(
    task_dir: Path = DEFAULT_TASK_DIR,
    *,
    enable_ocr: bool = True,
    ocr_languages: str = "rus+eng",
) -> list[CorpusDocument]:
    if not task_dir.exists():
        return []

    documents: list[CorpusDocument] = []
    for path in sorted(task_dir.glob("**/*")):
        if not path.is_file():
            continue
        if is_lfs_pointer(path):
            continue
        try:
            suffix = path.suffix.lower()
            if suffix == ".docx":
                documents.extend(_load_docx(path))
            elif suffix == ".pdf":
                documents.extend(_load_pdf(path))
            elif suffix == ".xlsx":
                documents.extend(_load_xlsx(path))
            elif suffix == ".png":
                documents.extend(
                    _load_png(path, enable_ocr=enable_ocr, ocr_languages=ocr_languages)
                )
        except (BadZipFile, OSError, ValueError) as exc:
            documents.append(
                CorpusDocument(
                    text=f"Источник {path.name} не удалось разобрать: {type(exc).__name__}: {exc}",
                    metadata=_metadata_for_path(path, _short_id("ERR", path), path.stem, "parse error"),
                    path=path,
                )
            )

    return documents


def load_corpus(
    include_task_data: bool = True,
    task_dir: Path = DEFAULT_TASK_DIR,
    *,
    enable_ocr: bool = True,
    ocr_languages: str = "rus+eng",
) -> list[CorpusDocument]:
    documents = load_demo_corpus()
    if include_task_data:
        documents.extend(
            load_task_corpus(task_dir, enable_ocr=enable_ocr, ocr_languages=ocr_languages)
        )
    return documents
